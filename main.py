import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
from tkcalendar import DateEntry
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import json
from datetime import datetime

# ===================== 配置与常量 =====================
CONFIG_FILE = "config.json"  # 基础信息配置
BUDGET_DATA_FILE = "budget_data.json"  # 预算数据持久化文件（一次导入后续复用）
EXCEL_SHEETS = ["施工项目（Sheet1）", "材料项目（Sheet2）"]
MAX_IMG_WIDTH = Inches(4)
MAX_IMG_HEIGHT = Inches(3)


class HomeAndEnterpriseTool:
    def __init__(self, root):
        self.root = root
        self.root.title("家集客项目预算与文档生成系统")
        self.root.geometry("1400x900")
        self.root.minsize(1300, 850)

        # 核心数据存储
        self.budget_data = []  # 整合后的预算项目
        self.total_amount = 0.0
        self.base_info = {}
        self.word_app_template = None
        self.word_review_template = None
        self.image_paths = []

        # ========== 修复：提前初始化status_var，解决属性不存在问题 ==========
        self.status_var = tk.StringVar(value="✅ 初始化中...")

        # 加载配置与预算数据（优先加载本地持久化数据）
        self.load_config()
        self.load_budget_data()  # 新增：加载本地持久化数据
        if not self.budget_data:  # 无本地数据时才导入Excel
            self.load_budget_excel()

        # 初始化GUI
        self.setup_style()
        self.setup_ui()

    # ===================== 数据持久化（新增核心功能）=====================
    def save_budget_data(self):
        """将预算数据保存到本地JSON文件（持久化）"""
        try:
            with open(BUDGET_DATA_FILE, "w", encoding="utf-8") as f:
                json.dump(self.budget_data, f, ensure_ascii=False, indent=2)
            self.status_var.set("✅ 预算数据已保存到本地，后续启动无需重新导入！")
        except Exception as e:
            messagebox.showerror("数据保存失败", f"错误原因：{str(e)}")

    def load_budget_data(self):
        """从本地JSON文件加载预算数据（持久化）"""
        if os.path.exists(BUDGET_DATA_FILE):
            try:
                with open(BUDGET_DATA_FILE, "r", encoding="utf-8") as f:
                    self.budget_data = json.load(f)
                # 重新生成连续ID
                for idx, item in enumerate(self.budget_data):
                    item["id"] = idx + 1
                messagebox.showinfo("加载成功", f"从本地加载{len(self.budget_data)}个预算项目（无需重新导入Excel）")
            except Exception as e:
                messagebox.showwarning("本地数据加载失败", f"将重新导入Excel：{str(e)}")
                self.budget_data = []
        else:
            self.budget_data = []

    # ===================== 原有基础配置加载/保存 =====================
    def load_config(self):
        default_info = {
            "申请单位": "奇台县分公司",
            "申请人": "樊斌",
            "联系电话": "13909949883",
            "实施单位": "中移建设",
            "项目经理": "吴斌",
            "项目经理联系电话": "18899661100",
            "项目负责人": "樊斌"
        }
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                    self.base_info = json.load(f)
                for key, val in default_info.items():
                    if key not in self.base_info:
                        self.base_info[key] = val
            except Exception as e:
                messagebox.showwarning("配置加载失败", f"使用默认基础信息：{str(e)}")
                self.base_info = default_info
        else:
            self.base_info = default_info
            self.save_config()

    def save_config(self):
        try:
            with open(CONFIG_FILE, "w", encoding="utf-8") as f:
                json.dump(self.base_info, f, ensure_ascii=False, indent=2)
        except Exception as e:
            messagebox.showerror("配置保存失败", str(e))

    # ===================== 预算表加载（保留原有逻辑，新增保存持久化）=====================
    def load_budget_excel(self):
        file_path = filedialog.askopenfilename(
            title="选择家集客预算表",
            filetypes=[("Excel文件", "*.xlsx"), ("所有文件", "*.*")]
        )
        if not file_path:
            messagebox.showwarning("提示", "未选择预算表，应用将无法正常使用！")
            return

        try:
            sheet1 = pd.read_excel(file_path, sheet_name=0)
            if sheet1.empty:
                raise ValueError("Sheet1（施工项目）为空")
            sheet1_data = self.parse_sheet1(sheet1)

            sheet2 = pd.read_excel(file_path, sheet_name=1)
            if sheet2.empty:
                raise ValueError("Sheet2（材料项目）为空")
            sheet2_data = self.parse_sheet2(sheet2)

            self.budget_data = sheet1_data + sheet2_data
            for idx, item in enumerate(self.budget_data):
                item["id"] = idx + 1
            # 新增：保存到本地持久化文件
            self.save_budget_data()
            messagebox.showinfo("加载成功",
                                f"共加载{len(self.budget_data)}个项目（施工{len(sheet1_data)}个+材料{len(sheet2_data)}个）")
        except Exception as e:
            messagebox.showerror("预算表加载失败", f"错误原因：{str(e)}")

    def parse_sheet1(self, df):
        parsed = []
        df.columns = df.columns.str.strip()
        required_cols = ["类别", "折扣后（含税）37%/元", "数量"]
        missing_cols = [col for col in required_cols if col not in df.columns]
        if missing_cols:
            raise ValueError(f"Sheet1缺少必要列：{', '.join(missing_cols)}")

        for _, row in df.iterrows():
            project_name = str(row["类别"]).strip()
            if not project_name or project_name == "nan":
                continue
            unit_price = float(pd.to_numeric(row["折扣后（含税）37%/元"], errors="coerce")) if pd.notna(
                row["折扣后（含税）37%/元"]) else 0.0
            quantity = float(pd.to_numeric(row["数量"], errors="coerce")) if pd.notna(row["数量"]) else 0.0
            is_length_unit = "元/公里" in project_name
            parsed.append({
                "id": len(parsed) + 1,
                "category": "施工项目",
                "name": project_name,
                "unit": "公里" if is_length_unit else "个/户/处等",
                "unit_price": unit_price,
                "quantity": quantity,
                "total": quantity * unit_price,
                "is_length": is_length_unit
            })
        if not parsed:
            raise ValueError("Sheet1（施工项目）无有效数据")
        return parsed

    def parse_sheet2(self, df):
        parsed = []
        df.columns = df.columns.str.strip()
        required_cols = ["材料", "含税", "数量"]
        missing_cols = [col for col in required_cols if col not in df.columns]
        if missing_cols:
            raise ValueError(f"Sheet2缺少必要列：{', '.join(missing_cols)}")

        for _, row in df.iterrows():
            project_name = str(row["材料"]).strip()
            if not project_name or project_name == "nan":
                continue
            unit_price = float(pd.to_numeric(row["含税"], errors="coerce")) if pd.notna(row["含税"]) else 0.0
            quantity = float(pd.to_numeric(row["数量"], errors="coerce")) if pd.notna(row["数量"]) else 0.0
            parsed.append({
                "id": len(parsed) + 1,
                "category": "材料项目",
                "name": project_name,
                "unit": "个",
                "unit_price": unit_price,
                "quantity": quantity,
                "total": quantity * unit_price,
                "is_length": False
            })
        if not parsed:
            raise ValueError("Sheet2（材料项目）无有效数据")
        return parsed

    # ===================== 样式配置 =====================
    def setup_style(self):
        self.style = ttk.Style(self.root)
        self.style.theme_use("clam")
        self.style.configure("Custom.TLabelframe", font=("Arial", 10), foreground="#333")
        self.style.configure("Custom.TLabelframe.Label", font=("Arial", 10, "bold"))
        self.style.configure("Accent.TButton", font=("Arial", 10), background="#4A90E2", foreground="white", padding=4)
        self.style.configure("Generate.TButton", font=("Arial", 11, "bold"), background="#2196F3", foreground="white",
                             padding=6)
        self.style.configure("Treeview.Heading", font=("Arial", 9, "bold"), background="#E0E0E0")
        self.style.configure("Treeview", font=("Arial", 8), rowheight=22)
        self.style.map("Treeview", background=[("selected", "#81C784")])

    # ===================== GUI界面布局（新增增删改查、导出按钮）=====================
    def setup_ui(self):
        # 1. 基础信息设置区
        base_frame = ttk.LabelFrame(self.root, text="📝 基础信息设置（设置后自动复用）", style="Custom.TLabelframe")
        base_frame.pack(fill=tk.X, padx=15, pady=8)

        info_grid = [
            ("申请单位", "申请单位"), ("申请人", "申请人"),
            ("联系电话", "联系电话"), ("实施单位", "实施单位"),
            ("项目经理", "项目经理"), ("项目经理联系电话", "项目经理联系电话"),
            ("项目负责人", "项目负责人")
        ]
        for i, (label, key) in enumerate(info_grid):
            row = i // 2
            col = i % 2
            ttk.Label(base_frame, text=f"{label}：", font=("Arial", 9)).grid(row=row, column=col * 3, padx=5, pady=5,
                                                                            sticky=tk.W)
            entry = ttk.Entry(base_frame, width=30, font=("Arial", 9))
            entry.grid(row=row, column=col * 3 + 1, padx=5, pady=5)
            entry.insert(0, self.base_info.get(key, ""))
            entry.bind("<FocusOut>", lambda e, k=key, ent=entry: self.update_base_info(k, ent.get()))

        ttk.Button(base_frame, text="💾 保存基础信息", command=self.save_config, style="Accent.TButton").grid(row=4,
                                                                                                             column=0,
                                                                                                             columnspan=6,
                                                                                                             pady=8)

        # 2. 预算表编辑区（新增增删改查按钮）
        budget_frame = ttk.LabelFrame(self.root, text="💰 预算项目编辑（仅工程量>0计入统计）", style="Custom.TLabelframe")
        budget_frame.pack(fill=tk.BOTH, padx=15, pady=5, expand=True)

        # 新增：增删改查按钮组
        btn_frame = ttk.Frame(budget_frame)
        btn_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Button(btn_frame, text="➕ 新增施工项目", command=self.add_construction_project,
                   style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="➕ 新增材料项目", command=self.add_material_project, style="Accent.TButton").pack(
            side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="🗑️ 删除选中项目", command=self.delete_selected_project,
                   style="Accent.TButton").pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="✏️ 修改项目信息", command=self.edit_project_info, style="Accent.TButton").pack(
            side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="📤 导出工程量>0项目到Excel", command=self.export_budget_to_excel,
                   style="Accent.TButton").pack(side=tk.LEFT, padx=5)

        # 标签页
        notebook = ttk.Notebook(budget_frame)
        notebook.pack(fill=tk.BOTH, padx=5, pady=5, expand=True)

        # 施工项目标签页（滚动条）
        construction_tab = ttk.Frame(notebook)
        construction_canvas = tk.Canvas(construction_tab)
        construction_vscroll = ttk.Scrollbar(construction_tab, orient=tk.VERTICAL, command=construction_canvas.yview)
        construction_hscroll = ttk.Scrollbar(construction_tab, orient=tk.HORIZONTAL, command=construction_canvas.xview)
        construction_scrollable_frame = ttk.Frame(construction_canvas)

        construction_scrollable_frame.bind("<Configure>", lambda e: construction_canvas.configure(
            scrollregion=construction_canvas.bbox("all")))
        construction_canvas.create_window((0, 0), window=construction_scrollable_frame, anchor="nw")
        construction_canvas.configure(yscrollcommand=construction_vscroll.set, xscrollcommand=construction_hscroll.set)

        self.construction_tree = self.create_treeview(construction_scrollable_frame, "施工项目")
        self.construction_tree.pack(fill=tk.BOTH, padx=5, pady=5, expand=True)

        construction_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        construction_vscroll.pack(side=tk.RIGHT, fill=tk.Y)
        construction_hscroll.pack(side=tk.BOTTOM, fill=tk.X)
        notebook.add(construction_tab, text="施工项目")

        # 材料项目标签页（滚动条）
        material_tab = ttk.Frame(notebook)
        material_canvas = tk.Canvas(material_tab)
        material_vscroll = ttk.Scrollbar(material_tab, orient=tk.VERTICAL, command=material_canvas.yview)
        material_hscroll = ttk.Scrollbar(material_tab, orient=tk.HORIZONTAL, command=material_canvas.xview)
        material_scrollable_frame = ttk.Frame(material_canvas)

        material_scrollable_frame.bind("<Configure>",
                                       lambda e: material_canvas.configure(scrollregion=material_canvas.bbox("all")))
        material_canvas.create_window((0, 0), window=material_scrollable_frame, anchor="nw")
        material_canvas.configure(yscrollcommand=material_vscroll.set, xscrollcommand=material_hscroll.set)

        self.material_tree = self.create_treeview(material_scrollable_frame, "材料项目")
        self.material_tree.pack(fill=tk.BOTH, padx=5, pady=5, expand=True)

        material_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        material_vscroll.pack(side=tk.RIGHT, fill=tk.Y)
        material_hscroll.pack(side=tk.BOTTOM, fill=tk.X)
        notebook.add(material_tab, text="材料项目")

        # 总金额显示
        self.total_var = tk.StringVar(value=f"当前总金额：0.00元")
        ttk.Label(budget_frame, textvariable=self.total_var, font=("Arial", 10, "bold"), foreground="#E64A19").pack(
            pady=5)

        # 3. 项目核心信息区
        project_frame = ttk.LabelFrame(self.root, text="📋 项目核心信息", style="Custom.TLabelframe")
        project_frame.pack(fill=tk.X, padx=15, pady=5)

        ttk.Label(project_frame, text="项目名称：", font=("Arial", 9)).grid(row=0, column=0, padx=5, pady=8, sticky=tk.W)
        self.project_name_var = tk.StringVar(value="广电项目光猫安装、开通")
        ttk.Entry(project_frame, textvariable=self.project_name_var, width=50, font=("Arial", 9)).grid(row=0, column=1,
                                                                                                       padx=5, pady=8)

        ttk.Label(project_frame, text="项目日期：", font=("Arial", 9)).grid(row=0, column=2, padx=15, pady=8,
                                                                           sticky=tk.W)
        self.date_entry = DateEntry(project_frame, width=20, background="#4A90E2", foreground="white",
                                    date_pattern="yyyy年MM月dd日", font=("Arial", 9))
        self.date_entry.grid(row=0, column=3, padx=5, pady=8)

        ttk.Label(project_frame, text="计划实施周期：", font=("Arial", 9)).grid(row=0, column=4, padx=15, pady=8,
                                                                               sticky=tk.W)
        self.cycle_var = tk.StringVar(value="15天")
        ttk.Entry(project_frame, textvariable=self.cycle_var, width=15, font=("Arial", 9)).grid(row=0, column=5, padx=5,
                                                                                                pady=8)

        # 4. 模板与支撑文件区
        template_frame = ttk.LabelFrame(self.root, text="📄 模板与支撑文件", style="Custom.TLabelframe")
        template_frame.pack(fill=tk.X, padx=15, pady=5)

        ttk.Label(template_frame, text="申请表模板：", font=("Arial", 9)).grid(row=0, column=0, padx=5, pady=6,
                                                                              sticky=tk.W)
        self.app_template_var = tk.StringVar(value="未选择")
        ttk.Entry(template_frame, textvariable=self.app_template_var, state="readonly", width=45,
                  font=("Arial", 9)).grid(row=0, column=1, padx=5, pady=6)
        ttk.Button(template_frame, text="浏览", command=lambda: self.select_template("app"),
                   style="Accent.TButton").grid(row=0, column=2, padx=5, pady=6)

        ttk.Label(template_frame, text="会审单模板：", font=("Arial", 9)).grid(row=1, column=0, padx=5, pady=6,
                                                                              sticky=tk.W)
        self.review_template_var = tk.StringVar(value="未选择")
        ttk.Entry(template_frame, textvariable=self.review_template_var, state="readonly", width=45,
                  font=("Arial", 9)).grid(row=1, column=1, padx=5, pady=6)
        ttk.Button(template_frame, text="浏览", command=lambda: self.select_template("review"),
                   style="Accent.TButton").grid(row=1, column=2, padx=5, pady=6)

        ttk.Label(template_frame, text="支撑图片（最多12张，仅插入申请表）：", font=("Arial", 9)).grid(row=0, column=3,
                                                                                                   padx=15, pady=6,
                                                                                                   sticky=tk.W)
        self.image_count_var = tk.StringVar(value="0张")
        ttk.Label(template_frame, textvariable=self.image_count_var, font=("Arial", 9)).grid(row=0, column=4, padx=5,
                                                                                             pady=6)
        ttk.Button(template_frame, text="上传", command=self.upload_images, style="Accent.TButton").grid(row=0,
                                                                                                         column=5,
                                                                                                         padx=5, pady=6)
        ttk.Button(template_frame, text="清空", command=self.clear_images, style="Accent.TButton").grid(row=0, column=6,
                                                                                                        padx=5, pady=6)

        # 5. 生成按钮
        self.generate_btn = ttk.Button(
            self.root, text="🚀 生成申请表+会审单", command=self.generate_documents, style="Generate.TButton"
        )
        self.generate_btn.pack(pady=15)

        # 状态提示
        self.status_var.set("✅ 基础信息已加载，可编辑预算项目工程量（双击表格修改）")
        status_label = ttk.Label(self.root, textvariable=self.status_var, font=("Arial", 9), foreground="#2196F3")
        status_label.pack(pady=5)

        # 刷新表格数据
        self.refresh_treeviews()

    # ===================== 表格创建与刷新 =====================
    def create_treeview(self, parent, category):
        tree = ttk.Treeview(
            parent,
            columns=["id", "name", "unit_price", "quantity", "total"],
            show="headings",
            selectmode="browse"
        )
        tree.heading("id", text="序号")
        tree.heading("name", text="项目名称")
        tree.heading("unit_price", text="单价（元）")
        tree.heading("quantity", text="工程量")
        tree.heading("total", text="合计（元）")
        tree.column("id", width=60)
        tree.column("name", width=450)
        tree.column("unit_price", width=100)
        tree.column("quantity", width=100)
        tree.column("total", width=100)
        tree.bind("<Double-1>", self.edit_quantity)
        return tree

    def refresh_treeviews(self):
        for item in self.construction_tree.get_children():
            self.construction_tree.delete(item)
        for item in self.material_tree.get_children():
            self.material_tree.delete(item)

        if not self.budget_data:
            self.total_var.set(f"当前总金额：0.00元")
            return

        self.total_amount = 0.0

        for item in self.budget_data:
            total = float(item["quantity"]) * float(item["unit_price"])
            item["total"] = total
            self.total_amount += total

            quantity_str = f"{float(item['quantity']):.2f}" if item["quantity"] is not None else "0.00"
            unit_price_str = f"{float(item['unit_price']):.2f}" if item["unit_price"] is not None else "0.00"
            total_str = f"{total:.2f}"

            values = [
                item["id"],
                item["name"],
                unit_price_str,
                quantity_str,
                total_str
            ]
            if item["category"] == "施工项目":
                self.construction_tree.insert("", tk.END, values=values, tags=("construction",))
            else:
                self.material_tree.insert("", tk.END, values=values, tags=("material",))

        self.total_var.set(f"当前总金额：{self.total_amount:.2f}元")

    # ===================== 增删改查功能（新增核心功能）=====================
    def add_construction_project(self):
        """新增施工项目"""
        # 弹出对话框输入项目信息
        name = simpledialog.askstring("新增施工项目", "请输入项目名称：")
        if not name:
            return
        unit_price = simpledialog.askfloat("新增施工项目", "请输入单价（元）：", initialvalue=0.0)
        if unit_price is None:
            return
        unit = simpledialog.askstring("新增施工项目", "请输入单位（如：公里、个）：", initialvalue="公里")
        if not unit:
            unit = "公里"
        is_length = simpledialog.askyesno("新增施工项目", "是否为长度类项目（单位：公里）？")
        quantity = simpledialog.askfloat("新增施工项目", "请输入工程量：", initialvalue=0.0)
        if quantity is None:
            quantity = 0.0

        # 添加到数据列表
        new_id = len(self.budget_data) + 1
        self.budget_data.append({
            "id": new_id,
            "category": "施工项目",
            "name": name.strip(),
            "unit": unit.strip(),
            "unit_price": unit_price,
            "quantity": quantity,
            "total": unit_price * quantity,
            "is_length": is_length
        })
        # 保存到本地并刷新表格
        self.save_budget_data()
        self.refresh_treeviews()
        self.status_var.set(f"✅ 新增施工项目：{name}")

    def add_material_project(self):
        """新增材料项目"""
        name = simpledialog.askstring("新增材料项目", "请输入材料名称：")
        if not name:
            return
        unit_price = simpledialog.askfloat("新增材料项目", "请输入单价（元）：", initialvalue=0.0)
        if unit_price is None:
            return
        unit = simpledialog.askstring("新增材料项目", "请输入单位（如：个）：", initialvalue="个")
        if not unit:
            unit = "个"
        quantity = simpledialog.askfloat("新增材料项目", "请输入工程量：", initialvalue=0.0)
        if quantity is None:
            quantity = 0.0

        new_id = len(self.budget_data) + 1
        self.budget_data.append({
            "id": new_id,
            "category": "材料项目",
            "name": name.strip(),
            "unit": unit.strip(),
            "unit_price": unit_price,
            "quantity": quantity,
            "total": unit_price * quantity,
            "is_length": False
        })
        self.save_budget_data()
        self.refresh_treeviews()
        self.status_var.set(f"✅ 新增材料项目：{name}")

    def delete_selected_project(self):
        """删除选中的项目"""
        # 判断当前选中的标签页（施工/材料）
        selected_item = None
        current_tree = None
        if self.construction_tree.focus():
            selected_item = self.construction_tree.focus()
            current_tree = self.construction_tree
        elif self.material_tree.focus():
            selected_item = self.material_tree.focus()
            current_tree = self.material_tree

        if not selected_item:
            messagebox.showwarning("提示", "请先选中要删除的项目！")
            return

        # 获取项目ID
        item_values = current_tree.item(selected_item)["values"]
        if len(item_values) < 1:
            messagebox.showwarning("提示", "选中项目数据异常！")
            return
        project_id = int(item_values[0])

        # 删除项目
        self.budget_data = [item for item in self.budget_data if item["id"] != project_id]
        # 重新生成ID
        for idx, item in enumerate(self.budget_data):
            item["id"] = idx + 1
        # 保存并刷新
        self.save_budget_data()
        self.refresh_treeviews()
        self.status_var.set(f"✅ 删除项目ID：{project_id}")

    def edit_project_info(self):
        """修改项目信息"""
        selected_item = None
        current_tree = None
        if self.construction_tree.focus():
            selected_item = self.construction_tree.focus()
            current_tree = self.construction_tree
        elif self.material_tree.focus():
            selected_item = self.material_tree.focus()
            current_tree = self.material_tree

        if not selected_item:
            messagebox.showwarning("提示", "请先选中要修改的项目！")
            return

        item_values = current_tree.item(selected_item)["values"]
        if len(item_values) < 4:
            messagebox.showwarning("提示", "选中项目数据异常！")
            return
        project_id = int(item_values[0])
        current_name = item_values[1]
        current_unit_price = float(item_values[2])
        current_quantity = float(item_values[3])

        # 查找项目
        target_item = None
        for item in self.budget_data:
            if item["id"] == project_id:
                target_item = item
                break
        if not target_item:
            messagebox.showwarning("提示", "项目不存在！")
            return

        # 弹出对话框修改
        new_name = simpledialog.askstring("修改项目信息", "请输入新的项目名称：", initialvalue=current_name)
        if not new_name:
            return
        new_unit_price = simpledialog.askfloat("修改项目信息", "请输入新的单价（元）：", initialvalue=current_unit_price)
        if new_unit_price is None:
            return
        new_quantity = simpledialog.askfloat("修改项目信息", "请输入新的工程量：", initialvalue=current_quantity)
        if new_quantity is None:
            return
        if target_item["category"] == "施工项目":
            new_unit = simpledialog.askstring("修改项目信息", "请输入新的单位：", initialvalue=target_item["unit"])
            if not new_unit:
                new_unit = target_item["unit"]
            new_is_length = simpledialog.askyesno("修改项目信息", "是否为长度类项目？",
                                                  initialvalue=target_item["is_length"])
            target_item["unit"] = new_unit.strip()
            target_item["is_length"] = new_is_length
        else:
            new_unit = simpledialog.askstring("修改项目信息", "请输入新的单位：", initialvalue=target_item["unit"])
            if not new_unit:
                new_unit = target_item["unit"]
            target_item["unit"] = new_unit.strip()

        # 更新信息
        target_item["name"] = new_name.strip()
        target_item["unit_price"] = new_unit_price
        target_item["quantity"] = new_quantity
        target_item["total"] = new_unit_price * new_quantity

        # 保存并刷新
        self.save_budget_data()
        self.refresh_treeviews()
        self.status_var.set(f"✅ 修改项目ID：{project_id}")

    def edit_quantity(self, event):
        """双击修改工程量（保留原有功能）"""
        tree = event.widget
        focus_item = tree.focus()
        if not focus_item:
            return
        item_values = tree.item(focus_item)["values"]
        if len(item_values) < 4:
            messagebox.showwarning("提示", "选中行数据不完整！")
            return
        try:
            project_id = int(item_values[0])
            current_quantity = float(item_values[3]) if item_values[3] not in ["nan", ""] else 0.0
        except (ValueError, IndexError):
            messagebox.showwarning("提示", "工程量数据异常！")
            return

        new_quantity = simpledialog.askfloat(
            "修改工程量",
            f"项目：{item_values[1]}\n当前工程量：{current_quantity:.2f}\n请输入新工程量（数字）：",
            initialvalue=current_quantity
        )
        if new_quantity is None:
            return
        if new_quantity < 0:
            messagebox.showwarning("警告", "工程量不能为负数！")
            return

        for item in self.budget_data:
            if item.get("id") == project_id:
                item["quantity"] = float(new_quantity)
                item["total"] = float(new_quantity) * float(item["unit_price"])
                break

        self.save_budget_data()  # 新增：保存修改后的数据
        self.refresh_treeviews()
        self.status_var.set(f"✅ 已更新项目工程量：{item_values[1]} → {new_quantity:.2f}")

    # ===================== 导出功能（新增核心功能）=====================
    def export_budget_to_excel(self):
        """导出工程量>0的项目到Excel"""
        # 筛选工程量>0的项目
        export_data = [item for item in self.budget_data if item["quantity"] > 0]
        if not export_data:
            messagebox.showwarning("提示", "无工程量>0的项目可导出！")
            return

        # 构造DataFrame
        df = pd.DataFrame({
            "序号": [item["id"] for item in export_data],
            "类别": [item["category"] for item in export_data],
            "项目名称": [item["name"] for item in export_data],
            "单位": [item["unit"] for item in export_data],
            "单价（元）": [item["unit_price"] for item in export_data],
            "工程量": [item["quantity"] for item in export_data],
            "合计（元）": [item["total"] for item in export_data]
        })

        # 保存文件
        save_path = filedialog.asksaveasfilename(
            title="导出预算项目到Excel",
            defaultextension=".xlsx",
            filetypes=[("Excel文件", "*.xlsx"), ("所有文件", "*.*")],
            initialfile=f"工程量大于0的预算项目_{datetime.now().strftime('%Y%m%d')}.xlsx"
        )
        if save_path:
            try:
                df.to_excel(save_path, index=False)
                messagebox.showinfo("导出成功", f"共导出{len(export_data)}个项目到Excel！")
                self.status_var.set(f"✅ 导出成功：{len(export_data)}个工程量>0的项目")
            except Exception as e:
                messagebox.showerror("导出失败", f"错误原因：{str(e)}")

    # ===================== 模板选择与图片上传 =====================
    def select_template(self, template_type):
        path = filedialog.askopenfilename(
            title=f"选择{'申请表' if template_type == 'app' else '会审单'}模板",
            filetypes=[("Word文件", "*.docx"), ("所有文件", "*.*")]
        )
        if not path:
            return
        if template_type == "app":
            self.word_app_template = path
            self.app_template_var.set(os.path.basename(path))
        else:
            self.word_review_template = path
            self.review_template_var.set(os.path.basename(path))
        self.status_var.set(f"✅ 已选择{'申请表' if template_type == 'app' else '会审单'}模板：{os.path.basename(path)}")

    def upload_images(self):
        paths = filedialog.askopenfilenames(
            title="选择支撑图片（仅插入申请表）",
            filetypes=[("图片文件", "*.jpg;*.jpeg;*.png;*.bmp"), ("所有文件", "*.*")]
        )
        if paths:
            remaining = 12 - len(self.image_paths)
            if len(paths) > remaining:
                messagebox.showwarning("提示", f"最多上传12张图片，本次仅上传{remaining}张！")
                paths = paths[:remaining]
            self.image_paths.extend(paths)
            self.image_count_var.set(f"{len(self.image_paths)}张")
            self.status_var.set(f"✅ 新增上传{len(paths)}张图片，累计{len(self.image_paths)}张（仅插入申请表）")

    def clear_images(self):
        self.image_paths.clear()
        self.image_count_var.set("0张")
        self.status_var.set("✅ 已清空所有支撑图片")

    def update_base_info(self, key, value):
        self.base_info[key] = value.strip()
        self.status_var.set(f"✅ 已更新{key}：{self.base_info[key]}（需点击保存按钮生效）")

    # ===================== 工作量清单生成（保留完整项目名称）=====================
    def generate_work_list(self):
        """生成工作量及材料清单（保留完整项目名称，不截断）"""
        work_list = []

        # 收集所有工程量>0的项目（施工+材料）
        for item in self.budget_data:
            if item["quantity"] <= 0:
                continue
            quantity = float(item["quantity"])
            # 保留完整项目名称，不截断
            if item["is_length"]:
                item_str = f"{quantity:.2f}公里 {item['name']}"
            else:
                item_str = f"{quantity:.2f}{item['unit']} {item['name']}"

            work_list.append(item_str)

        # 确保清单不为空
        return "，".join(work_list) if work_list else "无有效项目"

    # ===================== 图片插入辅助方法 =====================
    def insert_images_to_cell(self, cell, image_paths):
        if not image_paths:
            return
        cell.text = ""
        for img_path in image_paths:
            try:
                para = cell.add_paragraph()
                run = para.add_run()
                img = run.add_picture(img_path, width=MAX_IMG_WIDTH, height=MAX_IMG_HEIGHT)
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except Exception as e:
                messagebox.warning("图片插入失败", f"图片{os.path.basename(img_path)}插入失败：{str(e)}")

    def find_cell_by_text(self, table, keyword_list):
        """
        在表格中查找包含任意关键词的单元格，返回(行索引, 列索引, 单元格)
        :param table: docx的Table对象
        :param keyword_list: 关键词列表（如["工作量及材料清单", "主要工作量及材料清单"]）
        :return: (row_idx, col_idx, cell) 或 (None, None, None)
        """
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                for keyword in keyword_list:
                    if keyword in cell_text:
                        return (row_idx, col_idx, cell)
        return (None, None, None)

    # ===================== Word文档生成（核心修改：维修项目名称右侧列填写）=====================
    def generate_documents(self):
        if not self.word_app_template or not self.word_review_template:
            messagebox.showwarning("提示", "请先选择申请表和会审单模板！")
            return
        if self.total_amount <= 0:
            messagebox.showwarning("提示", "无有效项目（请设置工程量>0的项目）！")
            return
        project_name = self.project_name_var.get().strip()
        if not project_name:
            messagebox.showwarning("提示", "请输入项目名称！")
            return
        project_date = self.date_entry.get()
        cycle = self.cycle_var.get().strip()
        if not cycle:
            messagebox.showwarning("提示", "请输入计划实施周期！")
            return

        try:
            work_list = self.generate_work_list()

            self.fill_application_form(project_name, project_date, cycle, work_list)
            self.fill_review_form(project_name, project_date, cycle, work_list)

            messagebox.showinfo("生成成功",
                                f"✅ 两个文档已生成完成！\n总金额：{self.total_amount:.2f}元\n申请表中已插入{len(self.image_paths)}张支撑图片")
            self.status_var.set(f"🎉 生成成功！总金额：{self.total_amount:.2f}元，申请表插入{len(self.image_paths)}张图片")
        except IndexError as e:
            messagebox.showerror("生成失败",
                                 f"错误原因：模板表格行列索引越界（你的模板表格行列数与代码不匹配）\n详细错误：{str(e)}")
            self.status_var.set(f"❌ 生成失败：模板表格索引越界")
        except Exception as e:
            messagebox.showerror("生成失败", f"错误原因：{str(e)}")
            self.status_var.set(f"❌ 生成失败：{str(e)}")

    def fill_application_form(self, project_name, project_date, cycle, work_list):
        """填充申请表（核心修改：维修项目名称右侧列填写，保留原文字）"""
        doc = Document(self.word_app_template)
        if not doc.tables:
            raise ValueError("申请表模板中未找到表格！")
        target_table = doc.tables[0]

        # ========== 原有基础信息填充（保留，移除原项目名称填充项）==========
        fill_items = [
            (0, 1, self.base_info["申请单位"], WD_PARAGRAPH_ALIGNMENT.LEFT),
            (0, 3, project_date, WD_PARAGRAPH_ALIGNMENT.CENTER),
            (0, 4, project_date, WD_PARAGRAPH_ALIGNMENT.CENTER),
            (0, 6, self.base_info["申请人"], WD_PARAGRAPH_ALIGNMENT.LEFT),
            (1, 6, self.base_info["联系电话"], WD_PARAGRAPH_ALIGNMENT.LEFT),
            (2, 1, cycle, WD_PARAGRAPH_ALIGNMENT.LEFT),
            (2, 3, f"{self.total_amount:.2f}元", WD_PARAGRAPH_ALIGNMENT.CENTER),
            (2, 4, f"{self.total_amount:.2f}元", WD_PARAGRAPH_ALIGNMENT.CENTER),
        ]

        for row_idx, col_idx, text, align in fill_items:
            try:
                cell = target_table.cell(row_idx, col_idx)
                cell.text = text
                for para in cell.paragraphs:
                    para.alignment = align
            except IndexError:
                raise IndexError(f"申请表表格缺少行{row_idx}列{col_idx}的单元格")

        # ========== 核心修改1：维修项目名称 - 右侧列填写 ==========
        # 查找“维修项目名称”/“项目名称”单元格
        name_row_idx, name_col_idx, _ = self.find_cell_by_text(target_table, ["维修项目名称", "项目名称"])
        # 未找到时的兜底逻辑（保留原有的行1列1等位置，避免失效）
        if name_row_idx is None:
            name_row_idx = 1
            name_col_idx = 1
            # 提示用户模板不规范
            messagebox.showwarning("提示", "申请表模板中未找到“维修项目名称/项目名称”单元格，使用默认位置填充")

        # 确定填充列：原列的下一列（右侧列）
        name_fill_col = name_col_idx + 1
        # 边界处理：若超出列数，使用最后一列
        if name_fill_col >= len(target_table.columns):
            name_fill_col = len(target_table.columns) - 1

        # 填写项目名称（保留原单元格文字）
        name_fill_cell = target_table.cell(name_row_idx, name_fill_col)
        name_fill_cell.text = project_name
        # 设置格式
        for para in name_fill_cell.paragraphs:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in para.runs:
                run.font.size = Pt(10)

        # ========== 核心修改2：工作量及材料清单 - 右侧列填写（保留原逻辑）==========
        # 查找包含“工作量及材料清单”或“主要工作量及材料清单”的单元格
        list_row_idx, list_col_idx, _ = self.find_cell_by_text(target_table,
                                                               ["工作量及材料清单", "主要工作量及材料清单"])

        # 若未找到，扩大关键词范围（兼容原逻辑）
        if list_row_idx is None:
            list_row_idx, list_col_idx, _ = self.find_cell_by_text(target_table, ["工作量", "清单"])

        # 若仍未找到，默认使用倒数第3行第0列
        if list_row_idx is None:
            list_row_idx = max(0, len(target_table.rows) - 3)
            list_col_idx = 0

        # 确定填充列：原列的下一列（右侧列）
        fill_col_idx = list_col_idx + 1
        # 边界处理：若下一列超出表格列数，使用最后一列
        if fill_col_idx >= len(target_table.columns):
            fill_col_idx = len(target_table.columns) - 1

        # 获取填充单元格，填写内容（保留原单元格文字）
        fill_cell = target_table.cell(list_row_idx, fill_col_idx)
        fill_cell.text = work_list
        # 设置字体和对齐
        for para in fill_cell.paragraphs:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in para.runs:
                run.font.size = Pt(9)

        # ========== 图片插入逻辑保持不变 ==========
        keyword = "其他需求支撑文件"
        row_idx, col_idx, _ = self.find_cell_by_text(target_table, [keyword])
        if row_idx is not None and col_idx is not None:
            target_col_idx = col_idx + 1 if (col_idx + 1) < len(target_table.columns) else len(target_table.columns) - 1
            try:
                target_cell = target_table.cell(row_idx, target_col_idx)
                self.insert_images_to_cell(target_cell, self.image_paths)
            except IndexError:
                target_cell = target_table.cell(row_idx, len(target_table.columns) - 1)
                self.insert_images_to_cell(target_cell, self.image_paths)
        else:
            target_cell = target_table.cell(max(0, len(target_table.rows) - 2), len(target_table.columns) - 1)
            self.insert_images_to_cell(target_cell, self.image_paths)
            messagebox.showwarning("提示", "申请表模板中未找到“其他需求支撑文件”单元格，图片已插入到表格默认位置")

        save_path = filedialog.asksaveasfilename(
            title="保存申请表",
            defaultextension=".docx",
            filetypes=[("Word文件", "*.docx")],
            initialfile=f"{project_name}_申请表_{project_date.replace('年', '').replace('月', '').replace('日', '')}.docx"
        )
        if save_path:
            doc.save(save_path)

    def fill_review_form(self, project_name, project_date, cycle, work_list):
        """填充会审单（核心修改：维修项目名称右侧列填写，保留原文字）"""
        doc = Document(self.word_review_template)
        if not doc.tables:
            raise ValueError("会审单模板中未找到表格！")
        target_table = doc.tables[0]

        # ========== 原有基础信息填充（保留，移除原项目名称填充项）==========
        fill_items = [
            (1, 1, f"{self.total_amount:.2f}元", WD_PARAGRAPH_ALIGNMENT.CENTER),
            (1, 5, project_date, WD_PARAGRAPH_ALIGNMENT.CENTER),
            (1, 9, cycle, WD_PARAGRAPH_ALIGNMENT.CENTER),
            (2, 1, self.base_info["项目负责人"], WD_PARAGRAPH_ALIGNMENT.CENTER),
            (2, 5, self.base_info["联系电话"], WD_PARAGRAPH_ALIGNMENT.CENTER),
            (3, 1, self.base_info["实施单位"], WD_PARAGRAPH_ALIGNMENT.CENTER),
            (3, 5, self.base_info["项目经理"], WD_PARAGRAPH_ALIGNMENT.CENTER),
            (3, 9, self.base_info["项目经理联系电话"], WD_PARAGRAPH_ALIGNMENT.CENTER),
        ]

        for row_idx, col_idx, text, align in fill_items:
            try:
                if row_idx == 1 and col_idx == 1:
                    for c in range(1, 4):
                        cell = target_table.cell(row_idx, c)
                        cell.text = text
                        for para in cell.paragraphs:
                            para.alignment = align
                elif row_idx == 1 and col_idx == 5:
                    for c in range(5, 8):
                        cell = target_table.cell(row_idx, c)
                        cell.text = text
                        for para in cell.paragraphs:
                            para.alignment = align
                elif row_idx == 1 and col_idx == 9:
                    for c in range(9, min(11, len(target_table.columns))):
                        cell = target_table.cell(row_idx, c)
                        cell.text = text
                        for para in cell.paragraphs:
                            para.alignment = align
                else:
                    cell = target_table.cell(row_idx, col_idx)
                    cell.text = text
                    for para in cell.paragraphs:
                        para.alignment = align
            except IndexError:
                raise IndexError(f"会审单表格缺少行{row_idx}列{col_idx}的单元格")

        # ========== 核心修改1：维修项目名称 - 右侧列填写 ==========
        # 查找“维修项目名称”/“项目名称”单元格
        name_row_idx, name_col_idx, _ = self.find_cell_by_text(target_table, ["维修项目名称", "项目名称"])
        # 未找到时的兜底逻辑（保留原有的行0列1位置）
        if name_row_idx is None:
            name_row_idx = 0
            name_col_idx = 1
            messagebox.showwarning("提示", "会审单模板中未找到“维修项目名称/项目名称”单元格，使用默认位置填充")

        # 确定填充列：原列的下一列（右侧列）
        name_fill_col = name_col_idx + 1
        if name_fill_col >= len(target_table.columns):
            name_fill_col = len(target_table.columns) - 1

        # 填写项目名称（保留原单元格文字）
        name_fill_cell = target_table.cell(name_row_idx, name_fill_col)
        name_fill_cell.text = project_name
        # 若原逻辑是整行填充，可扩展为多个列（可选）
        # 例如：填充整行右侧列
        # for c in range(name_fill_col, len(target_table.columns)):
        #     target_table.cell(name_row_idx, c).text = project_name
        #     for para in target_table.cell(name_row_idx, c).paragraphs:
        #         para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for para in name_fill_cell.paragraphs:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in para.runs:
                run.font.size = Pt(10)

        # ========== 核心修改2：主要工作量及材料清单 - 右侧列填写（保留原逻辑）==========
        # 查找包含清单关键词的单元格
        list_row_idx, list_col_idx, _ = self.find_cell_by_text(target_table,
                                                               ["主要工作量及材料清单", "工作量及材料清单"])
        # 若未找到，扩大关键词范围
        if list_row_idx is None:
            list_row_idx, list_col_idx, _ = self.find_cell_by_text(target_table, ["工作量", "清单"])
        # 若仍未找到，默认使用倒数第2行第0列
        if list_row_idx is None:
            list_row_idx = max(0, len(target_table.rows) - 2)
            list_col_idx = 0

        # 确定清单填充列：右侧列
        list_fill_col = list_col_idx + 1
        if list_fill_col >= len(target_table.columns):
            list_fill_col = len(target_table.columns) - 1

        # 填写清单内容（带前缀）
        work_list_with_prefix = f"工作量：{work_list}"
        list_fill_cell = target_table.cell(list_row_idx, list_fill_col)
        list_fill_cell.text = work_list_with_prefix
        for para in list_fill_cell.paragraphs:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in para.runs:
                run.font.size = Pt(9)

        # ========== 核心修改3：施工方实施计划 - 右侧列填写（保留原逻辑）==========
        # 查找包含“施工方实施计划”的单元格
        plan_row_idx, plan_col_idx, _ = self.find_cell_by_text(target_table, ["施工方实施计划"])
        # 若未找到，默认在清单行的下一行，与清单同列
        if plan_row_idx is None:
            plan_row_idx = list_row_idx + 1
            plan_col_idx = list_col_idx

        # 确定计划填充列：右侧列
        plan_fill_col = plan_col_idx + 1
        if plan_fill_col >= len(target_table.columns):
            plan_fill_col = len(target_table.columns) - 1

        # 填写实施计划内容
        plan_text = f"我方计划安排1辆车2人在{cycle}完成施工。"
        plan_fill_cell = target_table.cell(plan_row_idx, plan_fill_col)
        plan_fill_cell.text = plan_text
        for para in plan_fill_cell.paragraphs:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in para.runs:
                run.font.size = Pt(9)

        save_path = filedialog.asksaveasfilename(
            title="保存会审单",
            defaultextension=".docx",
            filetypes=[("Word文件", "*.docx")],
            initialfile=f"{project_name}_会审单_{project_date.replace('年', '').replace('月', '').replace('日', '')}.docx"
        )
        if save_path:
            doc.save(save_path)


if __name__ == "__main__":
    root = tk.Tk()
    app = HomeAndEnterpriseTool(root)
    root.mainloop()